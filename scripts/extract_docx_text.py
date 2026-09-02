#!/usr/bin/env python3
"""Download one or more Word (.docx) documents and extract their text.

Less common than PDF for term dates, but it does happen (school office
staff often draft these in Word before ever converting to PDF, and
sometimes just link the .docx directly). Mirrors extract_pdf_text.py:
pure-Python (python-docx), no system dependencies, no OCR.

Old binary .doc (pre-2007 format) is NOT supported - python-docx only
reads the modern .docx XML format. If a .doc link is encountered, this
will report a clear error rather than silently producing garbage, so the
calling skill can fall back to asking the user to paste the text.

Usage:
    python3 extract_docx_text.py --url <docx_url> [--url <docx_url> ...]
"""
from __future__ import annotations

import argparse
import io
import json

import requests
from docx import Document

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
)
TIMEOUT = 30


def extract_one(url: str) -> dict:
    try:
        headers = {"User-Agent": USER_AGENT}
        response = requests.get(url, headers=headers, timeout=TIMEOUT)
        response.raise_for_status()
    except requests.RequestException as exc:
        return {"url": url, "error": True, "message": f"Could not download this document: {exc}"}

    if url.lower().split("?")[0].endswith(".doc"):
        return {
            "url": url,
            "error": True,
            "message": "This is an old-format .doc file, which isn't supported - only modern "
                       ".docx. Ask the user to open it and paste the text, or enter dates manually.",
        }

    try:
        document = Document(io.BytesIO(response.content))
    except Exception as exc:
        return {"url": url, "error": True, "message": f"Could not read this as a Word document: {exc}"}

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    extracted_text = "\n".join(parts).strip()
    return {
        "url": url,
        "extracted_text": extracted_text,
        "char_count": len(extracted_text),
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--url", action="append", required=True, dest="urls")
    args = parser.parse_args()

    sources = [extract_one(url) for url in args.urls]
    print(json.dumps({"sources": sources}, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
